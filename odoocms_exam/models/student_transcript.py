from odoo import models, fields, api
from odoo.exceptions import ValidationError, UserError
import pdb

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Row, CT_Tc
from docx.shared import Inches, Pt, Cm

import logging
import os
import tempfile
import subprocess
import sys
import json
import requests
from odoo import fields, models, _, api

from odoo.http import content_disposition, Controller, request, route
import re

_logger = logging.getLogger(__name__)

endpoint = 'http://198.38.85.147:5000/image'
endpoint2 = 'http://198.38.85.147:5000/files/transcript.pdf'


def preventDocumentBreak(document):
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)  # Append in the new tag


def set_table_width(table, data):
    i = 0
    for w in data:
        set_column_width(table.columns[i], Cm(data[i]))
        i += 1


def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)


def row_border(row, t=True, l=True, b=True, r=True):
    cells = row._tr.getchildren()
    for cell in cells:
        if isinstance(cell, CT_Tc):
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')

            top = OxmlElement('w:top')
            if t:
                top.set(qn('w:val'), 'single')
            else:
                top.set(qn('w:val'), 'nil')

            left = OxmlElement('w:left')
            if l:
                left.set(qn('w:val'), 'single')
            else:
                left.set(qn('w:val'), 'nil')

            bottom = OxmlElement('w:bottom')
            if b:
                bottom.set(qn('w:val'), 'single')
            else:
                bottom.set(qn('w:val'), 'nil')

            right = OxmlElement('w:right')
            if r:
                right.set(qn('w:val'), 'single')
            else:
                right.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)


def transcript_border(table):
    tbl = table._tbl
    rows = tbl.getchildren()
    lastRow = len(table.rows) - 1

    for row in rows:
        if isinstance(row, CT_Row):
            cells = row.getchildren()
            
            if row.tr_idx == 0:  # Top Row
                for cell in cells:
                    if isinstance(cell, CT_Tc):
                        tcPr = cell.tcPr
                        tcBorders = OxmlElement('w:tcBorders')

                        left = OxmlElement('w:left')
                        left.set(qn('w:val'), 'nil')

                        right = OxmlElement('w:right')
                        right.set(qn('w:val'), 'nil')

                        top = OxmlElement('w:top')
                        top.set(qn('w:val'), 'nil')

                        tcBorders.append(left)
                        tcBorders.append(right)
                        tcBorders.append(top)
                        tcPr.append(tcBorders)

            elif row.tr_idx == lastRow:  # Bottom Row
                j = 0
                for cell in cells:
                    if isinstance(cell, CT_Tc):
                        tcPr = cell.tcPr
                        tcBorders = OxmlElement('w:tcBorders')

                        top = OxmlElement('w:top')
                        top.set(qn('w:val'), 'nil')

                        if j > 0 and j < 3:  # Top Row
                            left = OxmlElement('w:left')
                            left.set(qn('w:val'), 'nil')
                            right = OxmlElement('w:right')
                            right.set(qn('w:val'), 'nil')
                            tcBorders.append(left)
                            tcBorders.append(right)

                        tcBorders.append(top)
                        tcPr.append(tcBorders)

                        j += 1
            else:
                j = 0
                for cell in cells:
                    if isinstance(cell, CT_Tc):
                        tcPr = cell.tcPr
                        tcBorders = OxmlElement('w:tcBorders')

                        top = OxmlElement('w:top')
                        top.set(qn('w:val'), 'nil')
                        bottom = OxmlElement('w:bottom')
                        bottom.set(qn('w:val'), 'nil')

                        if j > 0 and j < 3:  # Top Row
                            left = OxmlElement('w:left')
                            left.set(qn('w:val'), 'nil')
                            right = OxmlElement('w:right')
                            right.set(qn('w:val'), 'nil')
                            tcBorders.append(left)
                            tcBorders.append(right)

                        tcBorders.append(top)
                        tcBorders.append(bottom)
                        tcPr.append(tcBorders)

                        j += 1


def add_row(table, data, fontSize=7):
    row_cells = table.add_row().cells
    i = 0
    for val in data:
        tb_cell = row_cells[i].paragraphs[0]
        tb_cell.paragraph_format.space_before = Pt(0)
        tb_cell.paragraph_format.space_after = Pt(0)
        tb_cell_run = tb_cell.add_run()
        tb_cell_run.add_text(str(data[i]))
        if i == 2:
            tb_cell_run.font.size = Pt(fontSize - 2)
        else:
            tb_cell_run.font.size = Pt(fontSize)
        i += 1


class OdooCMSStudent(models.Model):
    _inherit = 'odoocms.student'

    def _download_transcript(self, model, report_type, report_ref, download=False):

        report_sudo = self.env.ref(report_ref).sudo()
        if not isinstance(report_sudo, type(self.env['ir.actions.report'])):
            raise UserError(_("%s is not the reference of a reports") % report_ref)

        method_name = 'render_qweb_doc'
        report = getattr(report_sudo, method_name)([model.id], data={'report_type': report_type})
        reporthttpheaders = [
            ('Content-Type', 'application/pdf' if report_type == 'docp' else 'text/html'),
            ('Content-Length', len(report)),
        ]

        if report_type == 'docp' and download:
            filename = "%s.pdf" % (re.sub('\W+', '-', "Transcript"))
            reporthttpheaders.append(('Content-Disposition', content_disposition(filename)))

        return request.make_response(report, headers=reporthttpheaders)

    def gen_transcript(self):
        document = Document()
        section = document.sections[-1]
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        tblw = [1.4, 4.9, 0.9, 0.9, 1.0]

        obj_styles = document.styles
        obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(8)
        obj_font.name = 'Times New Roman'

        obj_charstyle = obj_styles.add_style('CommentsStyle2', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(10)
        obj_font.name = 'Times New Roman'

        header = document.sections[0].header
        paragraph = header.add_paragraph()
        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        paragraph.add_run(self.env.user.company_id.name).bold = True

        paragraph = header.add_paragraph()
        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        paragraph.add_run('Registration No: ')
        paragraph.add_run(self.code or ' ').bold = True

        # paragraph = header.add_paragraph('Transcript', style='Intense Quote')
        # paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_table = header.add_table(1, 3, width=500)
        header_table.autofit = False

        A = header_table.cell(0, 0)
        pt = A.paragraphs[0]
        t = pt.text = ''
        pt.add_run("Student's Name: ", style='CommentsStyle').bold = True
        pt.add_run(self.name, style='CommentsStyle').bold = False
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(0)

        paragraph = A.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run("Program: ", style='CommentsStyle').bold = True
        paragraph.add_run(self.program_id.name, style='CommentsStyle').bold = False

        paragraph = A.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run("Plan: ", style='CommentsStyle').bold = True
        paragraph.add_run(self.program_id.name + ' Major', style='CommentsStyle').bold = False

        B = header_table.cell(0, 2)
        pt = B.paragraphs[0]
        t = pt.text = ''
        pt.add_run("Father's Name: ", style='CommentsStyle').bold = True
        pt.add_run(self.father_name, style='CommentsStyle').bold = False
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(0)

        cells = header_table.add_row().cells
        cells[0]._element.clear_content()
        table = cells[0].add_table(rows=0, cols=5)
        table.style = 'Table Grid'
        table.autofit = False

        set_table_width(table, tblw)
        add_row(table, ['Code', 'Title', ' ', 'CH', 'Grd'], fontSize=8)

        cells[2]._element.clear_content()
        table = cells[2].add_table(rows=0, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.style = 'Table Grid'
        table.autofit = False

        set_table_width(table, tblw)
        add_row(table, ['Code', 'Title', ' ', 'CH', 'Grd'], fontSize=8)

        set_table_width(header_table, [9.4, 1.0, 9.4])

        footer = document.sections[0].footer
        footer.is_linked_to_previous = False

        pt = footer.paragraphs[0]
        pt.add_run('"The Official Transcript carries the embossed stamp of the University"', style='CommentsStyle').bold = True
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(0)

        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('Transcript Prepared By: ---------------------------------------------', style='CommentsStyle').bold = False

        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('Transcript Checked By: ---------------------------------------------', style='CommentsStyle').bold = False

        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('Date of Issue: ' + str(fields.Date.today()), style='CommentsStyle').bold = False

        paragraph = footer.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('"Errors and Omissions are subject to Subsequent rectification"', style='CommentsStyle').bold = True
        paragraph.add_run("\t\t\tController of Examinations", style='CommentsStyle2').bold = True

        big_table = document.add_table(0, 1)
        big_table.autofit = False
        set_table_width(big_table, [9.5])

        student_terms =  self.term_ids + self.current_term_ids
        for term in student_terms:  # Student Semesters with Done state
            row = big_table.add_row()

            tag = row._tr
            child = OxmlElement('w:cantSplit')  # Create arbitrary tag
            tag.append(child)

            cells = row.cells
            cells[0]._element.clear_content()

            # label = cells[0].add_paragraph()
            # label.paragraph_format.keep_with_next = True
            # label.paragraph_format.space_before = Pt(0)
            # label.paragraph_format.space_after = Pt(0)

            table = cells[0].add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            a = table.cell(0, 0)
            b = table.cell(0, 4)
            A = a.merge(b)
            A.text = term.term_id.name
            A.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            # A.font.size = Pt(8)
            
            for course in term.student_course_ids:  # Student all Term Courses
                grade = (course.grade or ' ') if course.state == 'done' else ''
                rpt = course.repeat_code or ""
                if not course.include_in_cgpa:
                    rpt = '*'
                if course.course_id_1 or course.course_id_2 and not course.repeat_code:
                    rpt = '[RPT]'
                
                
                # Removed temp ***********************
                if course.include_in_transcript:
                    add_row(table, [
                        course.course_code or course.course_id.course_code or course.course_id.course_id.code,
                        course.course_name or course.course_id.course_name or course.course_id.course_id.name,
                        rpt,
                        course.credits,
                        grade
                    ])

                # course.course_id.course_id.code,

            set_table_width(table, tblw)
            transcript_border(table)

            row = table.add_row()
            row_border(row, b=False)
            a = table.cell(len(table.rows) - 1, 0)
            b = table.cell(len(table.rows) - 1, 4)
            A = a.merge(b)

            tb_cell_run = A.paragraphs[0].add_run()
            tb_cell_run.add_text("\tSCH: " + str(term.earned_credits))
            tb_cell_run.add_text("\tSGP: " + '{0:,.2f}'.format(term.grade_points))
            tb_cell_run.add_text("\tSGPA: " + '{0:,.2f}'.format(term.sgpa))
            tb_cell_run.font.size = Pt(7)

            row = table.add_row()
            # row_border(table.rows[len(table.rows)-1],t=False)
            row_border(row, t=False)
            a = table.cell(len(table.rows) - 1, 0)
            b = table.cell(len(table.rows) - 1, 4)
            A = a.merge(b)

            tb_cell_run = A.paragraphs[0].add_run()
            tb_cell_run.add_text("\tCCH: " + str(term.ecch))
            tb_cell_run.add_text("\tCGP: " + '{0:,.2f}'.format(term.cgp))
            tb_cell_run.add_text("\tCGPA: " + '{0:,.2f}'.format(term.cgpa))
            tb_cell_run.font.size = Pt(7)

            for row in table.rows:
                row.height = Cm(0.4)

            label = cells[0].paragraphs[0]
            label.paragraph_format.keep_with_next = True
            label.paragraph_format.space_before = Pt(0)
            label.paragraph_format.space_after = Pt(0)

        sectPr = document.sections[-1]._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')

        preventDocumentBreak(document)

        # document.save('demo.docx')

        temporary_files = []

        doc_report_fd, doc_report_path = tempfile.mkstemp(suffix='.docx', prefix='reports.tmp.')
        os.close(doc_report_fd)
        temporary_files.append(doc_report_path)

        pdf_report_fd, pdf_report_path = tempfile.mkstemp(suffix='.pdf', prefix=self.id_number.replace('/','_') or self.name or "" +'.tmp.')
        os.close(pdf_report_fd)
        # temporary_files.append(pdf_report_path)
        document.save(doc_report_path)

        # send to server
        # headers = {
        #     'Content-Type': 'multipart/form-data',
        # }
        #
        # response = requests.put(endpoint, files={'file': open(doc_report_path, 'rb')})
        # if response.status_code == 200:
        #     print(response.text)
        #
        # response = requests.get(endpoint2)
        # print(response.status_code)
        #
        # if response.status_code == 200:
        #     with open(pdf_report_path, 'wb') as out_file:  # change file name for PNG images
        #         out_file.write(response.content)

        pdf_converter = self.env['ir.config_parameter'].sudo().get_param('odoocms.pdf_converter')
        if not pdf_converter:
            raise UserError('Please configure Converter Path First.')
        try:

            wkhtmltopdf = [pdf_converter, "-f", "pdf", "-o", pdf_report_path, doc_report_path]
            process = subprocess.Popen(wkhtmltopdf, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            out, err = process.communicate()
        except:
            raise

        with open(pdf_report_path, 'rb') as pdf_document:
            pdf_content = pdf_document.read()

        # Manual cleanup of the temporary files
        for temporary_file in temporary_files:
            try:
                os.unlink(temporary_file)
            except (OSError, IOError):
                _logger.error('Error when trying to remove file %s' % temporary_file)

        return pdf_content,pdf_report_path

